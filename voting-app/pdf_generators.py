"""
Shared PDF generation functions for the FRCA Election App.

Extracted from app.py so that both Flask routes and CLI scripts
can generate PDFs without code duplication.
"""

import io
import math
import zipfile
from datetime import datetime

import qrcode
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm, cm
from reportlab.lib.colors import HexColor
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth

from name_formatting import shorten_to_fit

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

NAVY = HexColor("#1A3353")
GOLD = HexColor("#D4A843")

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _generate_qr_image(url, size=120):
    """Generate a QR code image as a ReportLab-compatible ImageReader object."""
    qr = qrcode.QRCode(
        version=None,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,
        border=1,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img_buf = io.BytesIO()
    img.save(img_buf, format="PNG")
    img_buf.seek(0)
    return ImageReader(img_buf)


# ---------------------------------------------------------------------------
# Warning strip (shared between front and back of dual-sided ballots)
# ---------------------------------------------------------------------------

_WARNING_STRIP_H = 6 * mm


def _draw_warning_strip(c, x, bottom_y, w, text):
    """Draw a grey warning strip with bold text at the bottom of a cell."""
    c.setFillColor(HexColor("#F0F0F0"))
    c.rect(x, bottom_y, w, _WARNING_STRIP_H, fill=1, stroke=0)
    c.setStrokeColor(HexColor("#000000"))
    c.setLineWidth(1.5)
    c.line(x, bottom_y + _WARNING_STRIP_H, x + w, bottom_y + _WARNING_STRIP_H)
    c.setLineWidth(1)
    c.setFont("Helvetica-Bold", 8)
    c.setFillColor(HexColor("#000000"))
    c.drawCentredString(x + w / 2, bottom_y + 2 * mm, text)


# ---------------------------------------------------------------------------
# Code slip drawing (used by both dual-sided and standalone code slips PDFs)
# ---------------------------------------------------------------------------


def _calc_code_slip_height(wifi_password):
    """Calculate the content height needed for a code slip."""
    h = 0
    h += 10 * mm  # header ("Vote Digitally" + rule + top padding)
    h += 17 * mm  # step 1 (WiFi) + gap to step 2
    h += 3.5 * mm  # password / "No password needed" line
    h += 32 * mm  # step 2 (QR row: label + QR/code side by side)
    h += 17 * mm  # dashed separator + step 3 (fallback)
    h += _WARNING_STRIP_H  # warning strip
    h += 2 * mm   # bottom padding
    return h


def draw_code_slip(c, x, top_y, w, cell_h, code, wifi_ssid, wifi_password,
                   base_url):
    """Draw a modern B&W-optimised code slip in the given cell.

    Layout: "Vote Digitally" header → Step 1 WiFi → Step 2 QR+code →
    Step 3 fallback URL (subdued) → warning strip.
    """
    cx = x + w / 2
    tx = x + 3 * mm
    bottom_y = top_y - cell_h
    y = top_y - 5 * mm

    # --- Solid border (rounded) ---
    c.setStrokeColor(HexColor("#000000"))
    c.setLineWidth(1.5)
    c.roundRect(x, bottom_y, w, cell_h, 2 * mm)
    c.setLineWidth(1)

    # --- Header: "Vote Digitally" + rule ---
    c.setFont("Helvetica-Bold", 13)
    c.setFillColor(HexColor("#000000"))
    c.drawCentredString(cx, y, "Vote Digitally")
    y -= 3.5 * mm
    c.setStrokeColor(HexColor("#000000"))
    c.setLineWidth(1.5)
    c.line(x + 3 * mm, y, x + w - 3 * mm, y)
    c.setLineWidth(1)
    y -= 6 * mm

    # --- Step number circle helper ---
    def _step_circle(sx, sy, num, active=True):
        r = 3.5 * mm
        cy = sy + 1 * mm  # centre circle on text baseline
        if active:
            c.setFillColor(HexColor("#000000"))
        else:
            c.setFillColor(HexColor("#DDDDDD"))
        c.circle(sx + r, cy, r, fill=1, stroke=0)
        if active:
            c.setFillColor(HexColor("#FFFFFF"))
        else:
            c.setFillColor(HexColor("#555555"))
        c.setFont("Helvetica-Bold", 9)
        c.drawCentredString(sx + r, cy - 1.2 * mm, str(num))

    # --- Step 1: Connect to WiFi ---
    _step_circle(tx, y, 1)
    label_x = tx + 9 * mm
    c.setFont("Helvetica", 10)
    c.setFillColor(HexColor("#777777"))
    c.drawString(label_x, y, "Connect to WiFi")
    ssid_x = label_x + c.stringWidth("Connect to WiFi ", "Helvetica", 10)
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(HexColor("#000000"))
    c.drawString(ssid_x, y, wifi_ssid)
    y -= 4 * mm
    c.setFont("Helvetica", 8)
    c.setFillColor(HexColor("#888888"))
    if wifi_password:
        c.drawString(label_x, y, f"Password: {wifi_password}")
    else:
        c.drawString(label_x, y, "No password needed")
    y -= 7 * mm

    # --- Step 2: Scan QR code ---
    _step_circle(tx, y, 2)
    c.setFont("Helvetica", 10)
    c.setFillColor(HexColor("#777777"))
    c.drawString(label_x, y, "Scan QR code with your phone camera")
    y -= 4 * mm

    # QR image only (voting code moved to step 3)
    qr_size = 24 * mm
    vote_url = f"{base_url}/v/{code}"
    qr_img = _generate_qr_image(vote_url)
    qr_x = label_x
    c.drawImage(qr_img, qr_x, y - qr_size, qr_size, qr_size)

    y -= qr_size + 2 * mm

    # --- OR divider + manual fallback (anchored above warning) ---
    y3 = bottom_y + _WARNING_STRIP_H + 14 * mm
    or_y = y3 + 7 * mm  # centre line for the OR pill

    # Draw line on each side of the "OR" pill
    or_text = "OR"
    c.setFont("Helvetica-Bold", 8)
    or_w = c.stringWidth(or_text, "Helvetica-Bold", 8) + 6 * mm  # pill width
    line_left = tx + 2 * mm
    line_right = x + w - 5 * mm
    pill_cx = x + w / 2
    pill_left = pill_cx - or_w / 2
    pill_right = pill_cx + or_w / 2

    c.setStrokeColor(HexColor("#CCCCCC"))
    c.setLineWidth(0.75)
    c.line(line_left, or_y, pill_left - 1 * mm, or_y)
    c.line(pill_right + 1 * mm, or_y, line_right, or_y)

    # Draw "OR" pill (rounded rect with dark fill)
    pill_h = 4.5 * mm
    c.setFillColor(HexColor("#333333"))
    c.roundRect(pill_left, or_y - pill_h / 2, or_w, pill_h,
                pill_h / 2, fill=1, stroke=0)
    c.setFillColor(HexColor("#FFFFFF"))
    c.drawCentredString(pill_cx, or_y - 1.3 * mm, or_text)

    # Fallback text: "Go to <url> and enter:"
    c.setFont("Helvetica", 10)
    c.setFillColor(HexColor("#555555"))
    c.drawString(tx, y3, "Go to")
    url_x = tx + c.stringWidth("Go to ", "Helvetica", 10)
    c.setFont("Helvetica-Bold", 10)
    c.setFillColor(HexColor("#000000"))
    c.drawString(url_x, y3, base_url)
    after_x = url_x + c.stringWidth(base_url + " ", "Helvetica-Bold", 10)
    c.setFont("Helvetica", 10)
    c.setFillColor(HexColor("#555555"))
    c.drawString(after_x, y3, "and enter:")
    y3 -= 5 * mm
    formatted_code = f"{code[:3]} {code[3:]}"
    c.setFont("Courier-Bold", 18)
    c.setFillColor(HexColor("#000000"))
    c.drawString(tx, y3, formatted_code)

    # --- Warning strip at bottom ---
    _draw_warning_strip(
        c, x, bottom_y, w,
        "\u26A0 Do not submit the paper ballot if you voted digitally"
    )


# ---------------------------------------------------------------------------
# Code Slips PDF
# ---------------------------------------------------------------------------


def generate_code_slips_pdf(codes, election_name, short_name, wifi_ssid,
                            wifi_password, base_url, is_demo=False):
    """Generate printable voting code cards — 6 per A4 page.

    Args:
        codes: list of voting code strings.
        election_name: name of the election.
        short_name: congregation short name (e.g. 'FRC Darling Downs').
        wifi_ssid: WiFi SSID to display.
        wifi_password: WiFi password to display (may be empty).
        base_url: voting base URL (e.g. 'http://192.168.8.100:5000').
        is_demo: if True, add demo watermark and header per page.

    Returns:
        BytesIO buffer containing the PDF.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    # Layout: 2 columns x N rows
    cols = 2
    margin = 8 * mm
    h_gap = 6 * mm
    v_gap = 6 * mm
    card_w = (width - 2 * margin - h_gap) / cols
    card_h = _calc_code_slip_height(wifi_password)
    rows_per_page = max(1, int((height - 2 * margin + v_gap) / (card_h + v_gap)))
    cards_per_page = cols * rows_per_page

    for page_start in range(0, len(codes), cards_per_page):
        page_codes = codes[page_start:page_start + cards_per_page]


        for i, code in enumerate(page_codes):
            row = i // cols
            col = i % cols

            card_x = margin + col * (card_w + h_gap)
            card_top_y = height - margin - row * (card_h + v_gap)

            draw_code_slip(c, card_x, card_top_y, card_w, card_h, code,
                           wifi_ssid, wifi_password, base_url)

        c.showPage()

    c.save()
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Counter Sheet PDF
# ---------------------------------------------------------------------------


def generate_counter_sheet_pdf(election_name, congregation_name, offices_data,
                               member_count=0, is_demo=False):
    """Generate a counter sheet PDF for paper ballot counting.

    Args:
        election_name: name of the election.
        congregation_name: full congregation name.
        offices_data: list of dicts, each with keys:
            'office': dict with 'name', 'max_selections', 'id'
            'candidates': list of dicts with 'name'
        member_count: number of members (for tick box count).
        is_demo: if True, add demo watermark and header per page.

    Returns:
        BytesIO buffer containing the PDF.
    """
    max_votes = max(member_count, 50)  # at least 50 boxes

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    margin = 8 * mm
    box_size = 3.5 * mm
    box_gap = 0.8 * mm
    group_gap = 2.5 * mm  # wider gap between groups of 5
    group_size = 5
    tally_start = 55 * mm
    total_col = width - 25 * mm
    # Calculate how many groups of 5 fit per row
    avail_width = total_col - tally_start - 5 * mm
    group_width = group_size * (box_size + box_gap) - box_gap + group_gap
    groups_per_row = int(avail_width / group_width)
    boxes_per_row = groups_per_row * group_size

    for office_item in offices_data:
        office = office_item["office"]
        candidates = office_item["candidates"]
        if not candidates:
            continue


        # Each office gets its own page(s)
        # Header
        c.setFillColor(NAVY)
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width / 2, height - 18 * mm, "Paper Ballot Counting Sheet")
        c.setFont("Helvetica", 9)
        c.setFillColor(HexColor("#666666"))
        c.drawCentredString(width / 2, height - 25 * mm,
                            f"{congregation_name} \u2014 {election_name}")

        c.setFont("Helvetica", 8)
        c.drawString(margin, height - 33 * mm,
                     "Counter: _________________________")
        c.drawString(width / 2, height - 33 * mm,
                     "Date: _______________")

        # Office title
        c.setFillColor(NAVY)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(margin, height - 42 * mm,
                     f"For {office['name']} (select {office['max_selections']})")

        y = height - 50 * mm

        for cand in candidates:
            # Calculate rows needed for this candidate
            row_height_est = box_size + box_gap + 1 * mm
            rows_needed = math.ceil(max_votes / boxes_per_row)
            candidate_height = 6 * mm + rows_needed * row_height_est + 8 * mm

            # Page break if needed
            if y - candidate_height < 25 * mm:
                c.showPage()
                y = height - 20 * mm

            # Candidate name -- shortened to fit between the left margin
            # and the start of the tally grid (with a small visual buffer).
            c.setFillColor(NAVY)
            c.setFont("Helvetica-Bold", 10)
            name_max_w = tally_start - margin - 4 * mm
            display_name = shorten_to_fit(
                cand["name"], name_max_w, "Helvetica-Bold", 10)
            c.drawString(margin, y, display_name)

            # Total box (right side)
            c.setStrokeColor(NAVY)
            c.setFont("Helvetica", 7)
            c.setFillColor(HexColor("#666666"))
            c.drawString(total_col, y, "Total")
            c.rect(total_col, y - rows_needed * row_height_est - 2 * mm,
                   12 * mm, rows_needed * row_height_est + 2 * mm)

            y -= 5 * mm

            # Draw tick box grid -- groups of 5, rows wrap on group boundaries
            row_height = box_size + box_gap + 1 * mm
            c.setStrokeColor(HexColor("#CCCCCC"))
            boxes_drawn = 0
            row = 0
            while boxes_drawn < max_votes:
                col_in_row = 0
                for g in range(groups_per_row):
                    if boxes_drawn >= max_votes:
                        break
                    gx = tally_start + g * group_width
                    by = y - row * row_height
                    for b in range(group_size):
                        if boxes_drawn >= max_votes:
                            break
                        bx = gx + b * (box_size + box_gap)
                        c.rect(bx, by, box_size, box_size)
                        boxes_drawn += 1
                        col_in_row += 1

                    # Group number label under the group
                    if boxes_drawn > 0 and boxes_drawn % group_size == 0:
                        c.setFont("Helvetica", 4.5)
                        c.setFillColor(HexColor("#BBBBBB"))
                        label_x = gx + (group_size * (box_size + box_gap) - box_gap) / 2
                        c.drawCentredString(label_x, by - 2.5 * mm, str(boxes_drawn))
                        c.setFillColor(NAVY)
                        c.setStrokeColor(HexColor("#CCCCCC"))

                row += 1

            actual_rows = row
            y -= actual_rows * row_height + 4 * mm

        # Footer
        c.setFont("Helvetica", 7)
        c.setFillColor(HexColor("#999999"))
        c.drawCentredString(width / 2, 12 * mm,
                            "Compare both counter sheets. Totals must match before entering into the app.")

        c.showPage()

    c.save()
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Paper Ballot PDF
# ---------------------------------------------------------------------------


def generate_paper_ballot_pdf(election_name, round_number, office_data,
                              member_count=0, is_demo=False):
    """Generate printable paper ballot forms.

    Args:
        election_name: name of the election.
        round_number: voting round number.
        office_data: list of dicts, each with keys:
            'office': dict with 'name', 'max_selections'
            'candidates': list of dicts with 'name'
        member_count: number of members (determines total ballots printed).
        is_demo: if True, add demo watermark and header per page.

    Returns:
        BytesIO buffer containing the PDF.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    n_offices = len(office_data)
    max_cands_in_office = (
        max(len(o["candidates"]) for o in office_data) if office_data else 0
    )

    # Single office uses the full ballot width; multiple offices split
    # left/right within the tile.
    two_office_cols = n_offices >= 2

    # Tile geometry (mm) — needed up front to compute the scale cap from
    # actual text widths.
    page_w_mm = 210
    margin_mm = 8
    col_gap_mm = 6
    sub_gap_mm = 3
    col_w_mm = (page_w_mm - 2 * margin_mm - col_gap_mm) / 2  # 94mm
    if two_office_cols:
        sub_w_mm = (col_w_mm - sub_gap_mm) / 2  # 45.5mm
    else:
        sub_w_mm = col_w_mm - 4  # ~90mm (single office, with side padding)

    # Compute the largest scale that keeps every candidate name and office
    # title inside its sub-column, AND fits 6 ballots per A4. The cand row
    # is: checkbox(3*scale mm) + 3mm gap + name. The office title is left-
    # aligned and must fit sub_w_mm. We aim for 6 ballots per A4 (3 rows
    # by 2 cols) so the printer can plan sheet count as ceil(ballots / 6).
    text_caps = [99.0]
    for item in office_data:
        office = item["office"]
        title = f"For {office['name']} (select {office['max_selections']})"
        title_w_mm = stringWidth(title, "Helvetica-Bold", 7) / mm
        if title_w_mm > 0:
            text_caps.append(sub_w_mm / title_w_mm)
        for cand in item["candidates"]:
            name_w_mm = stringWidth(cand["name"], "Helvetica", 7.5) / mm
            if name_w_mm > 0:
                # name_w * scale + box(3*scale) + 5mm padding <= sub_w
                text_caps.append((sub_w_mm - 5) / (name_w_mm + 3))
    text_fit_cap = min(text_caps)

    # Page-fit cap (target 6 ballots per A4 — 3 rows of 2 cols, ballot_h<=91mm)
    # ballot_h ≈ 14 + scale*(6.85 + 5*max_cands_in_office)
    # 14 + scale*(...)  <= 91  →  scale <= 77 / (6.85 + 5*max_cands)
    if max_cands_in_office > 0:
        page_fit_cap = 77 / (6.85 + 5 * max_cands_in_office)
    else:
        page_fit_cap = 99.0

    # Apply the smaller cap with a small safety margin (so text doesn't
    # touch the tile edge), and clamp to a reasonable range. 0.97 keeps
    # enough margin for descender/spacing while letting the font fill the
    # available width.
    scale = min(text_fit_cap, page_fit_cap) * 0.97
    scale = max(1.0, min(scale, 5.0))

    if two_office_cols:
        mid = (n_offices + 1) // 2
        left_offices = office_data[:mid]
        right_offices = office_data[mid:]
    else:
        left_offices = office_data
        right_offices = []

    # Heading row (title, optional round-2 warning) stays at base sizes
    # regardless of body scale — otherwise long election names overflow the
    # tile width into the adjacent ballot.
    title_pt = 9
    warning_pt = 6.5
    # Body elements scale together so checkboxes, names, and gaps grow
    # proportionally.
    box_mm = 3 * scale
    row_mm = 5 * scale
    office_header_mm = 4 * scale
    office_pad_mm = 1 * scale
    office_title_pt = 7 * scale
    cand_pt = 7.5 * scale
    # Fixed breathing room below the last candidate of an office. Without
    # this, drawing decrements `row_mm` after every candidate including the
    # last one, which (when row_mm is stretched) becomes a big empty band
    # at the bottom of the tile.
    tail_mm = 4

    def _col_body_height(offices):
        h = 0
        for item in offices:
            n = len(item["candidates"])
            if n == 0:
                h += (office_header_mm + office_pad_mm) * mm
            else:
                h += (office_header_mm + (n - 1) * row_mm + tail_mm + office_pad_mm) * mm
        return h

    body_height = max(_col_body_height(left_offices),
                      _col_body_height(right_offices) if right_offices else 0)

    # The body office-title ascender grows with scale, so reserve enough
    # clearance under the heading to keep "For Elder (select X)" from
    # crowding the election title.
    # Always reserve EXTRA breathing room above the body so the first
    # candidate doesn't sit right under the heading, regardless of how
    # densely packed the body is.
    EXTRA_ABOVE_BODY_MM = 5
    heading_to_body_gap = (
        max(4, 1.85 * scale + 3) + EXTRA_ABOVE_BODY_MM
    ) * mm
    # Header height = top inset + (round-2 warning slot) + heading-to-body gap.
    # Subtitle has been removed.
    header_height = 5 * mm + heading_to_body_gap
    if round_number > 1:
        header_height += 4 * mm
    padding = 6 * mm
    natural_ballot_h = header_height + body_height + padding
    # Standardize at 6 ballots per A4 so the printer can plan sheet count
    # as ceil(ballots / 6) regardless of slate shape. Fixed target is
    # ((usable_h + row_gap) / 3) - row_gap = 91mm.
    TARGET_BALLOT_H_MM = 91
    ballot_h = max(natural_ballot_h, TARGET_BALLOT_H_MM * mm)
    # First, increase office_header_mm (the gap between "For Elder (select N)"
    # and the first candidate) — more extra for fewer candidates, capped so
    # the body still fits inside the tile.
    natural_row_mm = row_mm  # save for cap below
    target_body_h_mm = (ballot_h - header_height - padding) / mm
    if max_cands_in_office > 0:
        natural_body_no_extras_mm = (
            office_header_mm
            + max(0, max_cands_in_office - 1) * natural_row_mm
            + tail_mm + office_pad_mm
        )
        max_office_extra_mm = max(0, target_body_h_mm - natural_body_no_extras_mm)
        # Sparse → bigger extra. N=2 → 8mm, N=4 → 6mm, N=6 → 4mm, N=8 → 3mm.
        desired_office_extra_mm = max(3, 10 - max_cands_in_office)
        office_header_mm += min(desired_office_extra_mm, max_office_extra_mm)

    # Then stretch the inter-candidate row spacing (only the gaps BETWEEN
    # candidates, not the trailing space) to fill what's left, capped at
    # 1.5x natural so we don't end up with comically large gaps. Any
    # residual slack gets centered evenly above and below.
    if max_cands_in_office > 1:
        candidate_room_mm = target_body_h_mm - office_header_mm - tail_mm - office_pad_mm
        stretched_row_mm = candidate_room_mm / (max_cands_in_office - 1)
        cap_row_mm = 1.5 * natural_row_mm
        row_mm = min(max(natural_row_mm, stretched_row_mm), cap_row_mm)
    # Recompute body_height with the chosen office_header and row spacing.
    body_height = max(_col_body_height(left_offices),
                      _col_body_height(right_offices) if right_offices else 0)
    # Centered residual slack — split evenly above and below the body.
    vcenter_offset = max(0, (ballot_h - (header_height + body_height + padding)) / 2)

    # Grid layout: 2 columns, as many rows as fit
    margin = 8 * mm
    col_gap = 6 * mm
    row_gap = 4 * mm
    col_w = (width - 2 * margin - col_gap) / 2
    usable_height = height - 2 * margin
    rows_per_page = max(1, int((usable_height + row_gap) / (ballot_h + row_gap)))
    ballots_per_page = rows_per_page * 2

    sub_gap = 3 * mm
    if two_office_cols:
        sub_w = (col_w - sub_gap) / 2
    else:
        sub_w = col_w - 4 * mm  # full width minus inset padding

    # Generate enough pages
    total_ballots = max(member_count + 10, 30) if member_count > 0 else 30

    ballot_index = 0
    while ballot_index < total_ballots:

        for slot in range(ballots_per_page):
            if ballot_index >= total_ballots:
                break

            col = slot % 2
            row = slot // 2

            x = margin + col * (col_w + col_gap)
            ballot_top = height - margin - row * (ballot_h + row_gap)

            # Dotted border
            c.setStrokeColor(HexColor("#CCCCCC"))
            c.setDash(2, 2)
            c.rect(x, ballot_top - ballot_h, col_w, ballot_h)
            c.setDash()

            # Content \u2014 heading row uses base font sizes regardless of body
            # scale. The heading-to-body gap is scale-aware so the body
            # office title doesn't crowd the heading at large scales.
            # vcenter_offset shifts everything down so the body sits in the
            # middle of the padded tile instead of being flushed to the top.
            cx = x + col_w / 2
            y = ballot_top - 5 * mm - vcenter_offset

            # Title
            c.setFillColor(NAVY)
            c.setFont("Helvetica-Bold", title_pt)
            c.drawCentredString(cx, y, election_name)

            if round_number > 1:
                y -= 4 * mm
                c.setFont("Helvetica-Bold", warning_pt)
                c.setFillColor(HexColor("#C0392B"))
                c.drawCentredString(cx, y, "Vote ONLY for candidates announced by the chairman")

            y -= heading_to_body_gap

            # Draw offices side by side (or single column for one office)
            body_y = y
            for ci, offices in enumerate([left_offices, right_offices]):
                if not offices:
                    continue
                ox = x + ci * (sub_w + sub_gap) + 2 * mm
                oy = body_y

                for item in offices:
                    office = item["office"]
                    candidates = item["candidates"]

                    c.setFillColor(NAVY)
                    c.setFont("Helvetica-Bold", office_title_pt)
                    c.drawString(ox, oy,
                                 f"For {office['name']} (select {office['max_selections']})")
                    oy -= office_header_mm * mm

                    name_max_w = sub_w - (box_mm + 5) * mm
                    last_idx = len(candidates) - 1
                    for ci_, cand in enumerate(candidates):
                        c.setStrokeColor(NAVY)
                        c.setFillColor(HexColor("#FFFFFF"))
                        c.rect(ox + 1 * mm, oy - 0.5 * mm, box_mm * mm, box_mm * mm)

                        c.setFillColor(NAVY)
                        c.setFont("Helvetica", cand_pt)
                        display_name = shorten_to_fit(
                            cand["name"], name_max_w, "Helvetica", cand_pt)
                        c.drawString(ox + (box_mm + 3) * mm, oy, display_name)
                        # Advance row_mm only BETWEEN candidates; after the
                        # last one, advance just `tail_mm` so we don't leave
                        # a row-sized empty band at the bottom of the office.
                        oy -= (row_mm if ci_ < last_idx else tail_mm) * mm

                    oy -= office_pad_mm * mm

            ballot_index += 1

        c.showPage()

    c.save()
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Results PDF
# ---------------------------------------------------------------------------


def generate_results_pdf(election_name, rounds_data, is_demo=False):
    """Export election results as a PDF.

    Args:
        election_name: name of the election.
        rounds_data: list of dicts, one per round, each with keys:
            'round_number': int
            'used_codes': int (digital votes cast)
            'offices': list of dicts, each with:
                'name': office name
                'candidates': list of dicts with 'name', 'digital', 'paper',
                              optional 'postal', and 'total'
        is_demo: if True, add demo watermark and header per page.

    Returns:
        BytesIO buffer containing the PDF.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=20 * mm, bottomMargin=20 * mm)
    styles = getSampleStyleSheet()
    elements = []

    # Title
    title_style = styles["Title"]
    title_style.textColor = NAVY
    elements.append(Paragraph(election_name, title_style))
    elements.append(Paragraph(
        f"Results as at {datetime.now().strftime('%d %B %Y %H:%M')}", styles["Normal"]
    ))
    elements.append(Spacer(1, 10 * mm))

    # Results for each round
    for round_data in rounds_data:
        round_num = round_data["round_number"]
        used_codes = round_data["used_codes"]

        postal_voter_count = round_data.get("postal_voter_count", 0)

        elements.append(Paragraph(f"Round {round_num}", styles["Heading2"]))
        summary = f"Digital votes cast: {used_codes}"
        if postal_voter_count:
            summary += f" &nbsp;|&nbsp; Postal voters: {postal_voter_count}"
        elements.append(Paragraph(summary, styles["Normal"]))
        elements.append(Spacer(1, 5 * mm))

        for office in round_data["offices"]:
            elements.append(Paragraph(f"{office['name']}", styles["Heading3"]))

            has_postal = any(c.get("postal", 0) > 0 for c in office["candidates"])

            candidates = office["candidates"]
            sum_digital = sum(c["digital"] for c in candidates)
            sum_paper = sum(c["paper"] for c in candidates)
            sum_postal = sum(c.get("postal", 0) for c in candidates)
            sum_total = sum(c["total"] for c in candidates)

            if has_postal:
                table_data = [["Candidate", "Digital", "Paper", "Postal", "Total"]]
                for cand in candidates:
                    table_data.append([
                        cand["name"],
                        str(cand["digital"]),
                        str(cand["paper"]),
                        str(cand.get("postal", 0)),
                        str(cand["total"]),
                    ])
                table_data.append([
                    "TOTAL", str(sum_digital), str(sum_paper),
                    str(sum_postal), str(sum_total),
                ])
                col_widths = [180, 65, 65, 65, 65]
            else:
                table_data = [["Candidate", "Digital", "Paper", "Total"]]
                for cand in candidates:
                    table_data.append([
                        cand["name"],
                        str(cand["digital"]),
                        str(cand["paper"]),
                        str(cand["total"]),
                    ])
                table_data.append([
                    "TOTAL", str(sum_digital), str(sum_paper), str(sum_total),
                ])
                col_widths = [200, 70, 70, 70]

            total_row = len(table_data) - 1
            table = Table(table_data, colWidths=col_widths)
            table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                ("TOPPADDING", (0, 0), (-1, 0), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -2),
                 [colors.white, HexColor("#F5F5F5")]),
                ("FONTNAME", (0, total_row), (-1, total_row), "Helvetica-Bold"),
                ("BACKGROUND", (0, total_row), (-1, total_row), HexColor("#E0E0E0")),
            ]))
            elements.append(table)
            elements.append(Spacer(1, 8 * mm))

    doc.build(elements)

    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Shared ballot card drawing (used by grid-based and printer exports)
# ---------------------------------------------------------------------------


def _draw_ballot_card(c, x, top_y, card_w, card_h, election_name,
                      left_offices, right_offices, sub_w, sub_gap):
    """Draw a single paper ballot card at the given position.

    Mirrors the dynamic-scale layout used by generate_paper_ballot_pdf:
    no subtitle, scale derived from actual text widths and available
    vertical room, sparse-aware office_header_mm, capped row stretching,
    and vertical centering of any residual slack. Single-office cards
    use the full card width (sub_w/sub_gap from the caller are ignored
    in that case).

    Args:
        c: ReportLab canvas.
        x: left edge of card.
        top_y: top edge of card.
        card_w: width of card.
        card_h: height of card.
        election_name: election title text.
        left_offices: list of office dicts for the left sub-column.
        right_offices: list of office dicts for the right sub-column (may
            be empty for single-office layouts).
        sub_w: width of each office sub-column (used only when both left
            and right have offices).
        sub_gap: gap between the two office sub-columns (same).
    """
    cx = x + card_w / 2
    bottom_y = top_y - card_h

    # Dashed cut border
    c.setStrokeColor(HexColor("#CCCCCC"))
    c.setDash(2, 2)
    c.rect(x, bottom_y, card_w, card_h)
    c.setDash()

    office_data_combined = list(left_offices) + list(right_offices)

    # Always draw the warning strip at the bottom; even an empty card
    # gets it.
    warning_text = "\u26A0 Do not submit this ballot if you voted digitally (see reverse)"
    if not office_data_combined:
        # Title only, then warning.
        c.setFillColor(HexColor("#000000"))
        c.setFont("Helvetica-Bold", 12)
        c.drawCentredString(cx, top_y - 5 * mm, election_name)
        _draw_warning_strip(c, x, bottom_y, card_w, warning_text)
        return

    n_offices = len(office_data_combined)
    two_office_cols = n_offices >= 2
    max_cands_in_office = max(len(o["candidates"]) for o in office_data_combined)

    # Single office uses the full card width; multi-office uses the
    # caller-supplied sub_w/sub_gap.
    if two_office_cols:
        sub_w_eff = sub_w
        sub_gap_eff = sub_gap
    else:
        sub_w_eff = card_w - 4 * mm
        sub_gap_eff = 0

    # Compute the largest scale that keeps every name and office title
    # inside its sub-column AND the body (heading + offices + warning)
    # fits the card height. Same approach as generate_paper_ballot_pdf.
    text_caps = [99.0]
    for item in office_data_combined:
        office = item["office"]
        title = f"For {office['name']} (select {office['max_selections']})"
        title_w_mm = stringWidth(title, "Helvetica-Bold", 7) / mm
        if title_w_mm > 0:
            text_caps.append((sub_w_eff / mm) / title_w_mm)
        for cand in item["candidates"]:
            name_w_mm = stringWidth(cand["name"], "Helvetica", 7.5) / mm
            if name_w_mm > 0:
                text_caps.append((sub_w_eff / mm - 5) / (name_w_mm + 3))
    text_fit_cap = min(text_caps)

    # Page-fit cap based on actual card_h. Solve for scale so:
    #   heading + body + bottom_padding + warning <= card_h
    # heading \u2248 5 + max(4, 1.85*scale + 3) + 5(extra)
    # body \u2248 scale * 5 * N + 4 (with tail_mm=4)
    # bottom_padding = 6, warning = _WARNING_STRIP_H_mm
    warn_h_mm = _WARNING_STRIP_H / mm
    fixed_h_mm = 5 + 5 + 3 + 4 + 6 + warn_h_mm  # baseline non-scale fixed parts
    available_for_scale_mm = (card_h / mm) - fixed_h_mm
    if available_for_scale_mm > 0 and max_cands_in_office > 0:
        page_fit_cap = available_for_scale_mm / (1.85 + 5 * max_cands_in_office)
    else:
        page_fit_cap = 1.0
    page_fit_cap = max(0.5, page_fit_cap)

    scale = min(text_fit_cap, page_fit_cap) * 0.97
    scale = max(1.0, min(scale, 5.0))

    # Per-element sizes (heading at base point sizes; body scales)
    title_pt = 12  # cards are larger than grid tiles, keep title at 12pt
    box_mm = 3 * scale
    natural_row_mm = 5 * scale
    row_mm = natural_row_mm
    natural_office_header_mm = 4 * scale
    office_header_mm = natural_office_header_mm
    office_pad_mm = 1 * scale
    office_title_pt = 7 * scale
    cand_pt = 7.5 * scale
    tail_mm = 4

    EXTRA_ABOVE_BODY_MM = 5
    heading_to_body_gap = (
        max(4, 1.85 * scale + 3) + EXTRA_ABOVE_BODY_MM
    ) * mm
    header_height = 5 * mm + heading_to_body_gap
    padding = 6 * mm
    warning_h = _WARNING_STRIP_H

    # Body region height = card_h - heading - padding - warning
    target_body_h_mm = (card_h - header_height - padding - warning_h) / mm

    # Sparse-aware office_header extra: more space above first cand for
    # fewer candidates. Capped so the body still fits.
    natural_body_no_extras_mm = (
        office_header_mm
        + max(0, max_cands_in_office - 1) * natural_row_mm
        + tail_mm + office_pad_mm
    )
    max_office_extra_mm = max(0, target_body_h_mm - natural_body_no_extras_mm)
    desired_office_extra_mm = max(3, 10 - max_cands_in_office)
    office_header_mm += min(desired_office_extra_mm, max_office_extra_mm)

    # Stretch row_mm (between candidates only), capped at 1.5x natural.
    if max_cands_in_office > 1:
        candidate_room_mm = target_body_h_mm - office_header_mm - tail_mm - office_pad_mm
        stretched_row_mm = candidate_room_mm / (max_cands_in_office - 1)
        cap_row_mm = 1.5 * natural_row_mm
        row_mm = min(max(natural_row_mm, stretched_row_mm), cap_row_mm)

    # Body height after choices \u2014 for centering
    def _col_body_h_mm(offices):
        h = 0
        for item in offices:
            n = len(item["candidates"])
            if n == 0:
                h += office_header_mm + office_pad_mm
            else:
                h += office_header_mm + (n - 1) * row_mm + tail_mm + office_pad_mm
        return h

    body_height = max(
        _col_body_h_mm(left_offices),
        _col_body_h_mm(right_offices) if right_offices else 0
    ) * mm
    vcenter_offset = max(0, (card_h - header_height - body_height - padding - warning_h) / 2)

    # ---- Draw heading ----
    cx = x + card_w / 2
    y = top_y - 5 * mm - vcenter_offset

    c.setFillColor(HexColor("#000000"))
    c.setFont("Helvetica-Bold", title_pt)
    c.drawCentredString(cx, y, election_name)
    y -= heading_to_body_gap

    # ---- Draw offices ----
    body_y = y
    cols_to_draw = [left_offices, right_offices] if two_office_cols else [left_offices]
    for ci, offices in enumerate(cols_to_draw):
        if not offices:
            continue
        if two_office_cols:
            ox = x + ci * (sub_w_eff + sub_gap_eff) + 2 * mm
        else:
            ox = x + 2 * mm
        oy = body_y

        for item in offices:
            office = item["office"]
            candidates = item["candidates"]

            c.setFillColor(HexColor("#000000"))
            c.setFont("Helvetica-Bold", office_title_pt)
            c.drawString(ox, oy,
                         f"For {office['name']} (select {office['max_selections']})")
            oy -= office_header_mm * mm

            name_max_w = sub_w_eff - (box_mm + 5) * mm
            last_idx = len(candidates) - 1
            for ci_, cand in enumerate(candidates):
                c.setStrokeColor(HexColor("#000000"))
                c.setFillColor(HexColor("#FFFFFF"))
                c.rect(ox + 1 * mm, oy - 0.5 * mm, box_mm * mm, box_mm * mm)

                c.setFillColor(HexColor("#000000"))
                c.setFont("Helvetica", cand_pt)
                display_name = shorten_to_fit(
                    cand["name"], name_max_w, "Helvetica", cand_pt)
                c.drawString(ox + (box_mm + 3) * mm, oy, display_name)
                oy -= (row_mm if ci_ < last_idx else tail_mm) * mm

            oy -= office_pad_mm * mm

    # Warning strip at bottom
    _draw_warning_strip(c, x, bottom_y, card_w, warning_text)


# ---------------------------------------------------------------------------
# Dual-Sided Ballots PDF (grid-based, for duplex printing)
# ---------------------------------------------------------------------------


def generate_dual_sided_ballots_pdf(election_name, short_name, round_number,
                                     office_data, codes, wifi_ssid,
                                     wifi_password, base_url,
                                     member_count=0, is_demo=False):
    """Generate grid-based dual-sided ballots PDF for duplex printing.

    Odd pages contain a grid of mini paper ballots. Even pages contain a grid
    of mini code slips at mirrored positions. When printed duplex long-edge on
    A4 and cut, each card has a paper ballot on one side and a code slip on
    the other.

    Args:
        election_name: e.g. "Office Bearer Election 2026"
        short_name: e.g. "FRC Darling Downs"
        round_number: current voting round (int)
        office_data: list of dicts, each with:
            'office': dict with 'name', 'max_selections', 'vacancies'
            'candidates': list of dicts with 'name'
        codes: list of plaintext 6-character code strings
        wifi_ssid: WiFi network name
        wifi_password: WiFi password (empty string for open network)
        base_url: e.g. "http://192.168.8.100:5000"
        member_count: number of members (limits cards printed to member_count + 10)
        is_demo: if True, add DEMO watermarks on every page

    Returns:
        BytesIO buffer containing the PDF.
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4

    margin = 8 * mm
    h_gap = 6 * mm
    v_gap = 6 * mm
    cols = 2
    col_w = (width - 2 * margin - h_gap) / cols

    # --- Side-by-side office layout ---
    mid = (len(office_data) + 1) // 2
    left_offices = office_data[:mid]
    right_offices = office_data[mid:]
    sub_gap = 3 * mm
    sub_w = (col_w - sub_gap) / 2

    # --- Calculate front (ballot) content height ---
    def _col_body_height(offices):
        h = 0
        for item in offices:
            h += 5.5 * mm + len(item["candidates"]) * 6 * mm + 1 * mm
        return h

    front_header_h = 5 * mm + 4.5 * mm + 4.5 * mm  # top pad + title + subtitle
    body_height = max(
        _col_body_height(left_offices),
        _col_body_height(right_offices) if right_offices else 0
    )
    front_h = front_header_h + body_height + _WARNING_STRIP_H + 2 * mm

    # --- Calculate back (code slip) content height ---
    back_h = _calc_code_slip_height(wifi_password)

    # --- cell_h = max of both sides ---
    cell_h = max(front_h, back_h)

    rows_per_page = max(1, int((height - 2 * margin + v_gap) / (cell_h + v_gap)))
    cards_per_page = cols * rows_per_page

    # Limit cards to member_count + 10 (spare), but never more than available codes
    total_cards = len(codes)

    # --- Helper: compute x, y for a grid slot ---
    def _slot_xy(slot, is_back):
        row = slot // cols
        if is_back:
            col = 1 - (slot % cols)  # mirror for long-edge duplex
        else:
            col = slot % cols
        x = margin + col * (col_w + h_gap)
        top_y = height - margin - row * (cell_h + v_gap)
        return x, top_y

    # --- Helper: draw a mini paper ballot in a cell (FRONT) ---
    def _draw_mini_ballot(slot):
        x, top_y = _slot_xy(slot, is_back=False)
        _draw_ballot_card(c, x, top_y, col_w, cell_h, election_name,
                          left_offices, right_offices, sub_w, sub_gap)

    # --- Generate pages ---
    for batch_start in range(0, len(codes), cards_per_page):
        batch_codes = codes[batch_start:batch_start + cards_per_page]

        # ODD PAGE: paper ballots (front)
        for slot in range(len(batch_codes)):
            _draw_mini_ballot(slot)
        c.showPage()

        # EVEN PAGE: code slips (back, columns mirrored)
        for slot, code_str in enumerate(batch_codes):
            sx, stop_y = _slot_xy(slot, is_back=True)
            draw_code_slip(c, sx, stop_y, col_w, cell_h, code_str,
                           wifi_ssid, wifi_password, base_url)
        c.showPage()

    c.save()
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Printer Pack — individual card-sized PDFs for professional printing
# ---------------------------------------------------------------------------


def _calc_card_dimensions(office_data, wifi_password):
    """Calculate card width, height, sub-column width and gap.

    Returns (col_w, cell_h, sub_w, sub_gap, left_offices, right_offices).
    Shared by grid-based and printer-pack generators.
    """
    h_gap = 6 * mm
    cols = 2
    width, _ = A4
    margin = 8 * mm
    col_w = (width - 2 * margin - h_gap) / cols

    mid = (len(office_data) + 1) // 2
    left_offices = office_data[:mid]
    right_offices = office_data[mid:]
    sub_gap = 3 * mm
    sub_w = (col_w - sub_gap) / 2

    def _col_body_height(offices):
        h = 0
        for item in offices:
            h += 5.5 * mm + len(item["candidates"]) * 6 * mm + 1 * mm
        return h

    front_header_h = 5 * mm + 4.5 * mm + 4.5 * mm
    body_height = max(
        _col_body_height(left_offices),
        _col_body_height(right_offices) if right_offices else 0
    )
    front_h = front_header_h + body_height + _WARNING_STRIP_H + 2 * mm
    back_h = _calc_code_slip_height(wifi_password)
    cell_h = max(front_h, back_h)

    return col_w, cell_h, sub_w, sub_gap, left_offices, right_offices


def generate_ballot_front_pdf(election_name, office_data, wifi_password,
                              is_demo=False):
    """Generate a single-page card-sized PDF with one paper ballot.

    The front is identical for all cards — the printer duplicates it.

    Returns:
        BytesIO buffer containing the PDF.
    """
    col_w, cell_h, sub_w, sub_gap, left, right = _calc_card_dimensions(
        office_data, wifi_password)

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(col_w, cell_h))

    _draw_ballot_card(c, 0, cell_h, col_w, cell_h, election_name,
                      left, right, sub_w, sub_gap)
    c.showPage()
    c.save()
    buf.seek(0)
    return buf


def generate_code_slips_back_pdf(codes, wifi_ssid, wifi_password, base_url,
                                 office_data, member_count=0, is_demo=False):
    """Generate card-sized PDF with one code slip per page.

    Each page has a unique voting code + QR. The printer cannot duplicate
    these — each page is different.

    Returns:
        BytesIO buffer containing the PDF.
    """
    col_w, cell_h, _, _, _, _ = _calc_card_dimensions(office_data,
                                                       wifi_password)

    total_cards = len(codes)

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(col_w, cell_h))

    for code_str in codes:
        draw_code_slip(c, 0, cell_h, col_w, cell_h, code_str,
                       wifi_ssid, wifi_password, base_url)
        c.showPage()

    c.save()
    buf.seek(0)
    return buf


def generate_cards_duplex_pdf(election_name, office_data, codes, wifi_ssid,
                              wifi_password, base_url, member_count=0,
                              is_demo=False):
    """Generate a card-sized PDF with interleaved front/back pages.

    Page layout: front, back, front, back, ... — 2N pages for N cards.
    Print duplex; no imposition setup needed. Each consecutive pair of
    pages produces one finished card.

    Returns:
        BytesIO buffer containing the PDF.
    """
    col_w, cell_h, sub_w, sub_gap, left, right = _calc_card_dimensions(
        office_data, wifi_password)

    total_cards = len(codes)

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=(col_w, cell_h))

    for code_str in codes:
        # Front (ballot)
        _draw_ballot_card(c, 0, cell_h, col_w, cell_h, election_name,
                          left, right, sub_w, sub_gap)
        c.showPage()

        # Back (unique code + QR)
        draw_code_slip(c, 0, cell_h, col_w, cell_h, code_str,
                       wifi_ssid, wifi_password, base_url)
        c.showPage()

    c.save()
    buf.seek(0)
    return buf


def generate_attendance_register_pdf(members, congregation_name,
                                     election_name=None, election_date=None):
    """Generate a printable attendance register PDF from a member list.

    Args:
        members: list of dicts with 'first_name' and 'last_name'.
        congregation_name: e.g. "Free Reformed Church of Darling Downs".
        election_name: optional election name for the header.
        election_date: optional election date string.

    Returns:
        BytesIO buffer containing the PDF.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=15 * mm,
                            bottomMargin=15 * mm, leftMargin=15 * mm,
                            rightMargin=15 * mm)
    elements = []
    styles = getSampleStyleSheet()

    title_style = styles["Title"]
    title_style.textColor = NAVY
    title_style.fontSize = 16
    elements.append(Paragraph("Attendance Register", title_style))
    elements.append(Paragraph(congregation_name, styles["Heading2"]))

    if election_name:
        details = election_name
        if election_date:
            details += f" \u2014 {election_date}"
        elements.append(Paragraph(details, styles["Normal"]))

    elements.append(Paragraph(
        "Article 4: All male communicant members present must sign this "
        "register.", styles["Normal"]))
    elements.append(Spacer(1, 5 * mm))

    table_data = [["#", "Name", "Signature"]]
    for i, member in enumerate(members, 1):
        name = f"{member['last_name']}, {member['first_name']}"
        table_data.append([str(i), name, ""])

    col_widths = [30, 200, 260]
    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("TOPPADDING", (0, 0), (-1, 0), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, HexColor("#F5F5F5")]),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ROWHEIGHT", (0, 1), (-1, -1), 36),
    ]))
    elements.append(table)

    doc.build(elements)
    buf.seek(0)
    return buf


def generate_av_instructions_pdf(election_name, wifi_ssid, wifi_password, base_url):
    """Generate a one-page A4 instruction sheet for the AV team.

    The election admin prints this and hands it to the person running the
    liturgy screen on election day. It explains how to put the /display page
    onto the screen from the AV booth PC.

    Args:
        election_name: name of the election (used as the page title).
        wifi_ssid: SSID of the election WiFi network.
        wifi_password: password for the election WiFi network.
        base_url: configured voting base URL (e.g. "http://church.vote").

    Returns:
        BytesIO buffer containing the PDF.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm,
                            bottomMargin=18 * mm, leftMargin=20 * mm,
                            rightMargin=20 * mm)
    elements = []
    styles = getSampleStyleSheet()

    # Title bar
    title_style = styles["Title"]
    title_style.textColor = NAVY
    title_style.fontSize = 26
    title_style.leading = 32
    elements.append(Paragraph("Notes for the AV team", title_style))
    subtitle_style = styles["Heading3"]
    subtitle_style.textColor = HexColor("#666666")
    subtitle_style.fontSize = 14
    subtitle_style.leading = 18
    elements.append(Paragraph(
        f"{election_name}: liturgy screen", subtitle_style))
    elements.append(Spacer(1, 8 * mm))

    # Step content styles
    h_style = getSampleStyleSheet()["Heading3"]
    h_style.textColor = NAVY
    h_style.fontSize = 18
    h_style.leading = 22
    h_style.spaceAfter = 2
    body_style = getSampleStyleSheet()["Normal"]
    body_style.fontSize = 14
    body_style.leading = 19

    display_url = f"{base_url.rstrip('/')}/display"
    password_text = wifi_password if wifi_password else "none"

    sections = [
        ("Connect to WiFi",
         f"Network: <b>{wifi_ssid}</b><br/>"
         f"Password: <b>{password_text}</b>"),
        ("Open in Chrome",
         f"<b>{display_url}</b>"),
        ("Press F11 for fullscreen",
         "Adjust zoom with <b>Ctrl +</b> / <b>Ctrl -</b>. "
         "<b>Ctrl 0</b> resets to 100%."),
        ("Leave on for the meeting",
         "The page updates automatically."),
        ("If anything doesn't display",
         f"Check you're still on the <b>{wifi_ssid}</b> WiFi, then press "
         "<b>F5</b> to refresh."),
    ]

    # Build a 2-column table: [numbered badge] [heading + body]
    badge_style = getSampleStyleSheet()["Normal"]
    badge_style.fontName = "Helvetica-Bold"
    badge_style.fontSize = 28
    badge_style.leading = 32
    badge_style.textColor = colors.white
    badge_style.alignment = 1  # center

    rows = []
    for idx, (heading, body) in enumerate(sections, 1):
        badge = Paragraph(str(idx), badge_style)
        content = [Paragraph(heading, h_style), Paragraph(body, body_style)]
        rows.append([badge, content])

    table = Table(rows, colWidths=[20 * mm, None])
    style_cmds = [
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 10),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
        ("LEFTPADDING", (1, 0), (1, -1), 12),
        # Navy badges in column 0
        ("BACKGROUND", (0, 0), (0, -1), NAVY),
        # Thin separator lines between rows
        ("LINEBELOW", (0, 0), (-1, -2), 0.5, HexColor("#DDDDDD")),
        # Outer card border
        ("BOX", (0, 0), (-1, -1), 1, HexColor("#CCCCCC")),
    ]
    table.setStyle(TableStyle(style_cmds))
    elements.append(table)

    doc.build(elements)
    buf.seek(0)
    return buf


def generate_printer_pack_zip(election_name, short_name, round_number,
                              office_data, codes, wifi_ssid, wifi_password,
                              base_url, congregation_name, members,
                              election_date=None, member_count=0,
                              is_demo=False):
    """Generate a ZIP containing all PDFs needed for professional printing.

    Contents (filenames prefixed so they sort in the order described
    in 0_INSTRUCTIONS.txt — read that first):

        0_INSTRUCTIONS.txt        — explanation of each file (read first)
        1_ballot_front.pdf        — card-sized paper ballot (duplicate this)
        2_code_slips_back.pdf     — N pages, card-sized unique code slips
        3_cards_duplex.pdf        — 2N pages, card-sized, interleaved front/back
        4_dual_sided_ballots.pdf  — grid layout for home duplex printing
        5_counter_sheet.pdf       — tally sheet for counting paper ballots
        6_attendance_register.pdf — sign-in sheet for election day
        7_av_instructions.pdf     — handout for the AV team

    Returns:
        BytesIO buffer containing the ZIP.
    """
    # 1. Ballot front (1 page, card-sized)
    front_buf = generate_ballot_front_pdf(
        election_name, office_data, wifi_password, is_demo=is_demo)

    # 2. Code slips back (N pages, card-sized)
    back_buf = generate_code_slips_back_pdf(
        codes, wifi_ssid, wifi_password, base_url, office_data,
        member_count=member_count, is_demo=is_demo)

    # 3. Cards duplex (interleaved front/back, card-sized, 2N pages)
    cards_duplex_buf = generate_cards_duplex_pdf(
        election_name, office_data, codes,
        wifi_ssid, wifi_password, base_url,
        member_count=member_count, is_demo=is_demo)

    # 4. Dual-sided grid layout (for home A4 printing fallback)
    dual_buf = generate_dual_sided_ballots_pdf(
        election_name, short_name, round_number, office_data, codes,
        wifi_ssid, wifi_password, base_url,
        member_count=member_count, is_demo=is_demo)

    # 4. Counter sheet
    counter_buf = generate_counter_sheet_pdf(
        election_name, congregation_name, office_data,
        member_count=member_count, is_demo=is_demo)

    # 5. Attendance register
    attendance_buf = generate_attendance_register_pdf(
        members, congregation_name,
        election_name=election_name, election_date=election_date)

    # 6. AV team instructions (handout for the liturgy screen operator)
    av_buf = generate_av_instructions_pdf(
        election_name, wifi_ssid, wifi_password, base_url)

    # 7. Instructions
    # One card per generated code. Multi-round elections generate codes
    # for all rounds at once (members x max_rounds), so the printer pack
    # produces all the cards needed for the whole election.
    total_cards = len(codes)
    total_cards_x2 = total_cards * 2
    instructions = f"""\
PRINTER PACK — {election_name}
{'=' * 60}

Thank you for printing the materials for our office bearer election.
Your work helps the congregation cast their votes on election day,
and we are grateful for your care and skill.

This ZIP contains everything needed to print materials for the
church office bearer election. Below is a description of each file.


CHOOSING A FORMAT
─────────────────
Three printing workflows are provided. Pick ONE based on your equipment:

  • Pro print shop with imposition software:    use #1 + #2
  • Card-size duplex printer, no imposition:    use #3
  • A4 home/office printer, no card media:      use #4


1. 1_ballot_front.pdf  (1 page)
   ─────────────────────────────
   The FRONT side of the voting card. Shows the election name,
   offices, and candidate checkboxes.

   This page is IDENTICAL for all cards. Your imposition software
   should duplicate it to produce {total_cards} copies, arranged on
   sheets for cutting. Card size: ~94 x 88 mm.

   Use this WITH file #2. Skip if using #3 or #4.


2. 2_code_slips_back.pdf  ({total_cards} pages)
   ─────────────────────────────
   The BACK side of the voting card. Each page has a UNIQUE voting
   code and QR code — one per card. Do NOT duplicate these pages.

   Page 1 pairs with copy 1 of the front, page 2 with copy 2, etc.
   Same card size as the front (~94 x 88 mm).

   Print these duplex with the front, matching page order:
   Front copy 1 + Back page 1, Front copy 2 + Back page 2, etc.

   Use this WITH file #1. Skip if using #3 or #4.


3. 3_cards_duplex.pdf  ({total_cards_x2} pages)
   ─────────────────────────────
   ALL-IN-ONE alternative to #1 + #2. Card-sized, with front and back
   interleaved on consecutive pages: page 1 is the front of card 1,
   page 2 is the back of card 1, page 3 is the front of card 2, etc.

   No imposition setup needed: send to a duplex printer that accepts
   card-size media and you get {total_cards} finished cards.

   Use INSTEAD of #1 + #2 if your printer can do card-size duplex.


4. 4_dual_sided_ballots.pdf
   ─────────────────────────────
   A4 FALLBACK for home/office printing without card media. Contains
   the same ballots in a 6-per-page grid layout, ready for duplex
   printing on A4 with long-edge binding. Cut along the dashed lines
   after printing.

   Use INSTEAD of #1 + #2 or #3 if you only have an A4 printer.


5. 5_counter_sheet.pdf
   ─────────────────────────────
   Tally sheet for counting paper ballots by hand. One page per
   office, with tick boxes for each candidate. Print on A4.


6. 6_attendance_register.pdf
   ─────────────────────────────
   Sign-in sheet listing all members. Each attendee signs next to
   their name upon arrival. Required per Article 4 of the church
   order. Print on A4.


7. 7_av_instructions.pdf  (1 page)
   ─────────────────────────────
   One-page handout for the AV team running the liturgy screen.
   The election admin gives this to the AV operator on the day.
   Print on A4. Not needed in bulk: 1-2 copies is enough.


PRINTING SUMMARY
{'=' * 60}

  File                        Size        Print     Yields              Paper
  ──────────────────────────  ──────────  ───────   ──────────────────  ──────
  1_ballot_front.pdf          Card-sized  x1        duplicated to {total_cards:<5}  Duplex (with #2)
  2_code_slips_back.pdf       Card-sized  x1        {total_cards} unique cards   Duplex (with #1)
  3_cards_duplex.pdf          Card-sized  x1        {total_cards} cards          Duplex (replaces #1+#2)
  4_dual_sided_ballots.pdf    A4          x1        {total_cards} cards (cut)    Duplex (A4 fallback)
  5_counter_sheet.pdf         A4          x2-3      tally sheet         Simplex
  6_attendance_register.pdf   A4          x1-2      sign-in sheet       Simplex
  7_av_instructions.pdf       A4          x1-2      AV handout          Simplex

For questions, contact the election administrator.
"""

    # Assemble ZIP. Filenames are prefixed 1_..7_ so they sort in the
    # order described in INSTRUCTIONS.txt when extracted.
    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("1_ballot_front.pdf", front_buf.getvalue())
        zf.writestr("2_code_slips_back.pdf", back_buf.getvalue())
        zf.writestr("3_cards_duplex.pdf", cards_duplex_buf.getvalue())
        zf.writestr("4_dual_sided_ballots.pdf", dual_buf.getvalue())
        zf.writestr("5_counter_sheet.pdf", counter_buf.getvalue())
        zf.writestr("6_attendance_register.pdf", attendance_buf.getvalue())
        zf.writestr("7_av_instructions.pdf", av_buf.getvalue())
        zf.writestr("0_INSTRUCTIONS.txt", instructions)

    zip_buf.seek(0)
    return zip_buf


# ---------------------------------------------------------------------------
# DOCX generation — Secretary's election minutes
# ---------------------------------------------------------------------------

def generate_minutes_docx(
    congregation_name,
    election_name,
    election_date,
    rounds_data,
    elected_summary,
    is_demo=False,
):
    """Generate a DOCX election minutes document for the secretary.

    Args:
        congregation_name: e.g. "Free Reformed Church of Darling Downs"
        election_name: e.g. "Office Bearer Election 2026"
        election_date: e.g. "4 October 2026"
        rounds_data: list of dicts per round, each with:
            'round_number': int
            'participants': int (in-person + postal)
            'in_person': int
            'postal_voter_count': int
            'used_codes': int
            'paper_ballot_count': int
            'total_ballots': int
            'offices': list of dicts, each with:
                'name': str
                'vacancies': int
                'max_selections': int
                'threshold_6a': float
                'threshold_6b': int
                'candidates': list of dicts with
                    'name', 'digital', 'paper', 'postal', 'total',
                    'elected' (bool)
        elected_summary: list of dicts with 'office' and 'names' (list of str)
        is_demo: if True, adds DEMO notice to header

    Returns:
        BytesIO buffer containing the DOCX file.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    navy = RGBColor(0x1A, 0x33, 0x53)

    # -- Helper: add a heading with navy colour --
    def _heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = navy
        return h

    # -- Helper: add a paragraph --
    def _para(text, bold=False, italic=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    # -- Helper: add a placeholder the secretary fills in --
    def _placeholder(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        return p

    # -- Helper: join names grammatically ("Brs A, B and C") --
    def _join_brs(names):
        names = list(names)
        if not names:
            return ""
        if len(names) == 1:
            return f"Br {names[0]}"
        return "Brs " + ", ".join(names[:-1]) + f" and {names[-1]}"

    def _vacancy_phrase(n):
        return f"{n} vacanc{'y' if n == 1 else 'ies'}"

    # =====================================================================
    # TITLE
    # =====================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(congregation_name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = navy

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Minutes of the Congregational Meeting "
        "for the Election of Office Bearers"
    )
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = navy

    if election_date:
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(election_date)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    # =====================================================================
    # 1. OPENING
    # =====================================================================
    _heading("1. Opening", level=2)
    _placeholder(
        "[The Chairman, Br [name]] opened the meeting at [time] and "
        "welcomed all present. The Congregation sang from [hymn]. He "
        "read from [scripture reference] and led in prayer."
    )

    # =====================================================================
    # 2. VOTING - narrative lead-in, then one sub-section per round
    # =====================================================================
    _heading("2. Voting", level=2)

    round1 = rounds_data[0] if rounds_data else None

    _para(
        "The Secretary, Br [name], read out Articles 4, 6 and 12 of the "
        "Rules for the Election of Office Bearers."
    )

    if round1:
        in_person = round1["in_person"]
        postal = round1["postal_voter_count"]
        if postal > 0:
            _para(
                f"The Chairman advised that a total of {in_person} male "
                "communicant members present had signed the attendance "
                f"register and that {postal} postal vote"
                f"{'' if postal == 1 else 's'} had been received. "
                "Brs [name] and [name] were appointed to assist with the "
                "collection and counting of votes."
            )
        else:
            _para(
                f"The Chairman advised that a total of {in_person} male "
                "communicant members present had signed the attendance "
                "register. Brs [name] and [name] were appointed to assist "
                "with the collection and counting of votes."
            )

    total_rounds = len(rounds_data)

    for rd_idx, rd in enumerate(rounds_data):
        round_num = rd["round_number"]
        is_last_round = (rd_idx == total_rounds - 1)
        _heading(f"2.{rd_idx + 1} Round {round_num}", level=3)

        # Narrative introduction to the round
        office_phrases = [
            f"the office of {o['name']} ({_vacancy_phrase(o['vacancies'])})"
            for o in rd["offices"]
        ]
        if round_num == 1:
            if len(office_phrases) == 1:
                intro = f"Voting was conducted for {office_phrases[0]}."
            elif len(office_phrases) == 2:
                intro = (
                    f"Voting was conducted for {office_phrases[0]} "
                    f"and {office_phrases[1]}."
                )
            else:
                intro = (
                    "Voting was conducted for "
                    + ", ".join(office_phrases[:-1])
                    + f", and {office_phrases[-1]}."
                )
        else:
            parts = []
            for o in rd["offices"]:
                names = [c["name"] for c in o["candidates"]]
                parts.append(
                    f"the office of {o['name']} "
                    f"({_vacancy_phrase(o['vacancies'])} remaining) "
                    f"between {_join_brs(names)}"
                )
            if len(parts) == 1:
                intro = f"A further ballot was conducted for {parts[0]}."
            else:
                intro = (
                    "A further ballot was conducted for "
                    + "; and for ".join(parts)
                    + "."
                )
        _para(intro)

        # Brief threshold mention per office
        for o in rd["offices"]:
            t6a = o.get("threshold_6a")
            t6b = o.get("threshold_6b")
            if t6a is not None and t6b is not None:
                _para(
                    f"For the office of {o['name']}, a candidate required "
                    f"more than {t6a:.2f} votes (Article 6a) and at least "
                    f"{t6b} votes (Article 6b) to be elected."
                )

        # Round-level ballot totals with digital / in-person / postal split
        rd_total = rd.get("total_ballots", 0)
        rd_digital = rd.get("used_codes", 0)
        rd_paper = rd.get("paper_ballot_count", 0)
        rd_postal = rd.get("postal_voter_count", 0)
        if rd_total > 0:
            def _pct(n):
                return f"{(100 * n / rd_total):.1f}%"
            if rd_postal > 0:
                _para(
                    f"A total of {rd_total} ballots were cast in this round: "
                    f"{rd_digital} digital ({_pct(rd_digital)}), "
                    f"{rd_paper} in-person ({_pct(rd_paper)}), "
                    f"and {rd_postal} postal ({_pct(rd_postal)})."
                )
            else:
                _para(
                    f"A total of {rd_total} ballots were cast in this round: "
                    f"{rd_digital} digital ({_pct(rd_digital)}) "
                    f"and {rd_paper} in-person ({_pct(rd_paper)})."
                )

        # Vote tables per office. The candidate breakdown is always shown
        # as Digital + In-person, with Postal added in round 1 when any
        # postal votes were received.
        for o in rd["offices"]:
            p = doc.add_paragraph()
            run = p.add_run(o["name"])
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)

            cands = o["candidates"]
            if not cands:
                _placeholder(f"[No candidates stood for {o['name']} this round.]")
                continue

            has_postal = round_num == 1 and any(
                c.get("postal", 0) > 0 for c in cands
            )

            if has_postal:
                headers = ["Candidate", "Digital", "In-person", "Postal", "Total"]
                col_widths = [Cm(6.0), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5)]
            else:
                headers = ["Candidate", "Digital", "In-person", "Total"]
                col_widths = [Cm(7.0), Cm(3.0), Cm(3.0), Cm(3.0)]
            col_count = len(headers)

            table = doc.add_table(rows=1 + len(cands), cols=col_count)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.allow_autofit = False
            for row in table.rows:
                for ci, width in enumerate(col_widths):
                    row.cells[ci].width = width

            for i, hdr in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = hdr
                if i > 0:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)

            for row_idx, cand in enumerate(cands):
                row = table.rows[row_idx + 1]
                row.cells[0].text = f"Br {cand['name']}"
                row.cells[1].text = str(cand.get("digital", 0))
                row.cells[2].text = str(cand.get("paper", 0))
                if has_postal:
                    row.cells[3].text = str(cand.get("postal", 0))
                    row.cells[4].text = str(cand["total"])
                else:
                    row.cells[3].text = str(cand["total"])
                for ci in range(1, col_count):
                    for paragraph in row.cells[ci].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for ci in range(col_count):
                    for paragraph in row.cells[ci].paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                if cand.get("elected"):
                    for ci in range(col_count):
                        for paragraph in row.cells[ci].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

        # Declaration sentence
        elected_clauses = []
        remaining_offices = []
        for o in rd["offices"]:
            elected_names = [c["name"] for c in o["candidates"] if c.get("elected")]
            unfilled = o["vacancies"] - len(elected_names)
            if elected_names:
                elected_clauses.append({
                    "office": o["name"],
                    "names": elected_names,
                })
            if unfilled > 0:
                remaining_offices.append((o["name"], unfilled))

        if elected_clauses:
            def _clause_text(c):
                return f"for the office of {c['office']}, {_join_brs(c['names'])}"
            if len(elected_clauses) == 1:
                names = elected_clauses[0]["names"]
                verb = "were" if len(names) > 1 else "was"
                _para(
                    f"The Chairman declared that {_join_brs(names)} {verb} "
                    f"elected for the office of {elected_clauses[0]['office']}."
                )
            else:
                joined = "; and ".join(_clause_text(c) for c in elected_clauses)
                _para(
                    "The Chairman declared that the following brothers "
                    f"were elected: {joined}."
                )
        else:
            _para(
                "The Chairman declared that no candidate was elected in "
                "this round, and a further ballot would be required."
            )

        # Outcome after this round
        if is_last_round:
            if remaining_offices:
                parts = [
                    f"{name} ({count} unfilled)"
                    for name, count in remaining_offices
                ]
                _para(
                    "The election for " + ", ".join(parts) +
                    " concluded without all vacancies being filled."
                )
            else:
                _para("All vacancies were now filled.")

            # Final result sentence (end-of-voting summary). Skip when
            # there was only one round — the chairman's declaration above
            # already lists the same brothers and the repetition reads
            # awkwardly. Multi-round elections still get the summary so
            # readers don't have to assemble it from each round's
            # declaration.
            if elected_summary and total_rounds > 1:
                clauses = []
                for item in elected_summary:
                    if item["names"]:
                        clauses.append(
                            f"for the office of {item['office']}, "
                            f"{_join_brs(item['names'])}"
                        )
                if clauses:
                    p = doc.add_paragraph()
                    run = p.add_run(
                        "The final result of the election is: "
                        + "; and ".join(clauses)
                        + "."
                    )
                    run.bold = True
        elif remaining_offices:
            parts = [
                f"the office of {name} ({count} vacanc{'y' if count == 1 else 'ies'} to fill)"
                for name, count in remaining_offices
            ]
            _para(
                "A further ballot was required for "
                + ", ".join(parts) + "."
            )

    # =====================================================================
    # 3. ARTICLE 12 - Objections
    # =====================================================================
    _heading("3. Objections (Article 12)", level=2)
    _placeholder(
        "The Chairman provided opportunity for any objections to be "
        "raised, noting that Article 12 of the Rules requires that any "
        "objections of a formal nature against procedure must be lodged "
        "at the meeting. [No objections were raised. / Record any "
        "objections here.]"
    )

    # =====================================================================
    # 4. CLOSING
    # =====================================================================
    _heading("4. Closing", level=2)
    _placeholder(
        "Br [name] led in prayer, and the Chairman closed the meeting "
        "at [time]."
    )

    # =====================================================================
    # SIGNATURE BLOCK
    # =====================================================================
    doc.add_paragraph()
    doc.add_paragraph()

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_table.rows[0].cells[0].text = "Chairman:"
    sig_table.rows[0].cells[1].text = "Secretary:"
    sig_table.rows[1].cells[0].text = "\n\n____________________________"
    sig_table.rows[1].cells[1].text = "\n\n____________________________"

    for row in sig_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    # =====================================================================
    # Footer note
    # =====================================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "This document was generated by the FRCA Election App. "
        "The secretary should verify all details and complete the "
        "sections marked in grey before filing in the minute book."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
